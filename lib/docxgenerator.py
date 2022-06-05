from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.shared import Pt, Inches

# noinspection HttpUrlsUsage
WNS_COLS_NUM = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"


class DocxGenerator:
    def __init__(self, filename):
        self.filename = filename
        self.p = None

    def __enter__(self):
        document = Document()
        style = document.styles["Normal"]
        font = style.font
        font.name = "Galatia sil"
        font.size = Pt(10.5)

        # Make the document two columns, most efficient way I found, but uses internal attributes of python-docx
        section = document.add_section(WD_SECTION_START.CONTINUOUS)
        # noinspection PyProtectedMember
        section._sectPr.xpath("./w:cols")[0].set(WNS_COLS_NUM, str(2))

        paragraph_format = document.styles['Normal'].paragraph_format
        paragraph_format.left_indent = Inches(0.25)
        paragraph_format.first_line_indent = Inches(-0.25)
        self.document = document

        self.new_paragraph()

        return self

    def new_paragraph(self):
        self.p = self.document.add_paragraph()

    def write(self, text, bold=False):
        out = self.p.add_run(text)
        if bold:
            out.bold = True

    def __exit__(self, _, __, ___):
        self.document.save(self.filename)
