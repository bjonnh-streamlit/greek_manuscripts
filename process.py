from cltk.alphabet.text_normalization import cltk_normalize
from cltk.lemmatize import GreekBackoffLemmatizer
from collections import defaultdict
# noinspection PyPackageRequirements
from docx import Document
# noinspection PyPackageRequirements
from docx.enum.section import WD_SECTION_START
# noinspection PyPackageRequirements
from docx.shared import Pt, Inches
from itertools import groupby
import re
import sys
import textract
from greek_accentuation.characters import base

FILE = "data/Galen Simpl Med 01 (Convert'd) Books 06-11.doc"
# Limits the output files size so they are easier to look through
DEBUG = False

# Sections are either enclosed in [] or starting with vol
regexp_section = re.compile(r"((^\[.*]$)|(^vol.*$))")
regexp_line = re.compile(r"([\d.]*)(\s?)(.*)")


def greek_word_basifier(word):
    """A function that helps sorting words by their base letters not diacritics ones
    We add the word at the end to have a correct ordering."""
    str_word = str(word)
    sort_key = "".join([base(ch) for ch in str_word]) + str_word
    return sort_key


class Reference:
    def __init__(self, section, subsection, line):
        self.section = section
        self.subsection = subsection
        self.line = line

    def nice_print(self):
        return f"{self.subsection} ({self.section}.{self.line})"


class Decoder:
    def __init__(self):
        self.section = ""
        self.subsection = ""
        self.line_counter = 1
        self.nlp = None

        # word to positions as list
        self.word_occurrences = defaultdict(list)

        # reversed index
        self.reversed_word_occurrences = defaultdict(list)

    def set_section(self, section):
        self.section = section.replace("[", "").replace("]", "")

    def current_reference(self):
        return Reference(self.section, self.subsection, self.line_counter)

    def process_text_line(self, text_line):
        words = text_line.split()
        for word in words:
            cleaned_word = word.replace(",", "").replace(".", "").replace(":", "").replace("·", "").strip()
            if cleaned_word != "":
                self.word_occurrences[cleaned_word].append(self.current_reference())
                self.reversed_word_occurrences[cleaned_word[::-1]].append(self.current_reference())

    def process_line(self, _line):
        clean_line = _line.strip()
        if clean_line == "":
            return

        line_match = regexp_line.match(_line)
        if regexp_section.match(clean_line):
            self.set_section(clean_line)
            self.line_counter = 0
        elif line_match:
            self.line_counter += 1
            subsection = line_match.group(1)
            if subsection != "":
                self.subsection = subsection
            self.process_text_line(line_match.group(3).strip())
        else:
            print(f"Can't handle {clean_line}")

    @staticmethod
    def nice_printer_references(content):
        """Prints the sections only when it is a new one."""
        output = ""
        for thing in groupby(content, lambda entry: entry.subsection):
            list_of_subref = [f"{ref.section}.{ref.line}" for ref in thing[1]]
            output += f"{thing[0]} ({' ‖ '.join(list_of_subref)})"
        return output

    def word_info(self, word, source):
        content = source[word]

        return f"({len(content)}) {self.nice_printer_references(content)}"

    def index(self, reverse=False):
        if reverse:
            source = self.reversed_word_occurrences
        else:
            source = self.word_occurrences

        for word in sorted(source.keys(), key=lambda kv: greek_word_basifier(kv)):
            yield word, decoder.word_info(word, source)

    def count(self):
        counts = {}

        for word in self.word_occurrences.keys():
            counts[word] = len(self.word_occurrences[word])
        # We have to use a little trick here as we want to reverse the numbers but not the words.
        return sorted(counts.items(), key=lambda kv: (sys.maxsize - kv[1], greek_word_basifier(kv[0])))

    def lemma(self, debug=False):
        """Produce the lemmatized sorted output"""
        lemmatizer = GreekBackoffLemmatizer()
        output = defaultdict(set)

        words = set()
        for word in self.word_occurrences.keys():
            words.add(cltk_normalize(word))
        if debug:
            print(f" Normalized, found {len(words)} unique words")
        for word in words:
            lemma = lemmatizer.lemmatize([word])[0]
            output[lemma[1]].add(word)

        output_sorted = [(entry[0], sorted(entry[1], key=lambda kv: greek_word_basifier(kv))) for entry in
                         sorted(output.items(), key=lambda kv: greek_word_basifier(kv[0]))]
        return output_sorted


# noinspection HttpUrlsUsage
WNS_COLS_NUM = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"


class DocxGenerator:
    def __init__(self, filename):
        self.filename = filename
        self.p = None

    def __enter__(self):
        document = Document()
        style = document.styles["Normal"]
        font = style.font
        font.name = "Galatia sil"
        font.size = Pt(10.5)

        # Make the document two columns, most efficient way I found, but uses internal attributes of python-docx
        section = document.add_section(WD_SECTION_START.CONTINUOUS)
        # noinspection PyProtectedMember
        section._sectPr.xpath("./w:cols")[0].set(WNS_COLS_NUM, str(2))

        paragraph_format = document.styles['Normal'].paragraph_format
        paragraph_format.left_indent = Inches(0.25)
        paragraph_format.first_line_indent = Inches(-0.25)
        self.document = document

        self.new_paragraph()

        return self

    def new_paragraph(self):
        self.p = self.document.add_paragraph()

    def write(self, text, bold=False):
        out = self.p.add_run(text)
        if bold:
            out.bold = True

    def __exit__(self, _, __, ___):
        self.document.save(self.filename)


full_text = textract.process(FILE).decode('utf-8')

decoder = Decoder()

count = 100

for line in full_text.splitlines():
    decoder.process_line(line)
    if DEBUG:
        count -= 1
        if count == 0:
            break

print("Generating direct index")
with DocxGenerator("out/index.docx") as generator:
    for i in decoder.index():
        generator.write(i[0], bold=True)
        generator.write(f": {i[1]}")
        generator.new_paragraph()

print("Generating inverted index")
with DocxGenerator("out/index_inverse.docx") as generator:
    for i in decoder.index(reverse=True):
        generator.write(i[0], bold=True)
        generator.write(f": {i[1]}")
        generator.new_paragraph()

print("Generating frequency list")
with DocxGenerator("out/frequency.docx") as generator:
    for i in decoder.count():
        generator.write(i[0], bold=True)
        generator.write(f": {i[1]}")
        generator.new_paragraph()

print("Generating lemma list")
with DocxGenerator("out/lemma.docx") as generator:
    for i in decoder.lemma(debug=True):
        generator.write(i[0], bold=True)
        generator.write(f": {', '.join(i[1])}")
        generator.new_paragraph()
